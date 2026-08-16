# Ce fichier va se charger de l'extraction : transformer le contenu du fichier chargé peu importe son type en texte
# Il va par exemple garantir que .py est simplement du texte et pas executer du code python;
# lambda/services/document_service.py


from pathlib import Path
import csv
import io
import json

from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".csv", ".txt", ".py", ".md", ".json",
}


def extract_pdf(file_bytes:bytes) -> str: 
    """
        Extrait le texte d'un fichier PDF.
    """

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""

        if text.strip():
            pages.append(f"===== PAGE {page_number} =====\n{text}")
    return "\n\n".join(pages)


def extract_docx(file_bytes:bytes) -> str:
    """
        Extrait les paragraphes et tableaux d'un fichier DOCX
    """

    document = Document(io.BytesIO(file_bytes))

    parts=[]

    #Paragraphes
    for paragraph in document.paragraphs:
        text=paragraph.text.strip()

        if text:
            parts.append(text)

    #Tableaux

    for table_index, table in enumerate(document.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ") for cell in row.cells
            ]

            rows.append(" | ".join(cells))

        if rows:
            parts.append(
                f"===== TABLEAU {table_index} =====\n" + "\n".join(rows)
            )

    return "\n\n".join(parts)


def extract_xlsx(file_bytes:bytes) -> str:
    """
        Extrait le contenu textuel de toutes les feuilles Excel.
    """
    workbook = load_workbook(
        filename=io.BytesIO(file_bytes),
        read_only=True,
        data_only=True,
    )

    sheets = []

    for worksheet in workbook.worksheets:
        rows =[]

        for row in worksheet.iter_rows(values_only=True):
            values = []

            for value in row:
                if value is None:
                    values.append("")
                else:
                    values.append(str(value))

            # Ignore les lignes totalement vides

            if any(value.strip() for value in values):
                rows.append(" | ".join(values))

        if rows:
            sheets.append(
                f"=====FEUILLE : {worksheet.title} =====\n" + "\n".join(rows)
            )

    workbook.close()
    return "\n\n".join(sheets)



def extract_csv(file_bytes:bytes) -> str:
    """
        Extrait un CSV en représentation textuelle.
    """

    decoded = file_bytes.decode("utf-8-sig", errors="replace")

    sample= decoded[:4096]

    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect= csv.excel

    reader = csv.reader(
        io.StringIO(decoded),
        dialect,
    )

    rows = []

    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def extract_text(file_bytes:bytes) -> str:
    """
        Extrait TXT, PY, MD, JSON.
    """
    return file_bytes.decode("utf-8", errors="replace",)


def extract_document(
        filename: str,
        file_bytes: bytes,) -> str:

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Format non supporté: {extension}"
        )

    if extension == ".pdf":
        return extract_pdf(file_bytes)

    if extension == ".docx":
        return extract_docx(file_bytes)

    if extension == ".xlsx":
        return extract_xlsx(file_bytes)

    if extension == ".csv":
        return extract_csv(file_bytes)

    if extension in {".txt", ".py", ".md",".json"}:
        return extract_text(file_bytes)

    raise ValueError(
        f"Format non supporté : {extension}"
    )